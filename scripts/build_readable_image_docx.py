from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


PAGE_W = 1275
PAGE_H = 1650
MARGIN_X = 88
TOP = 72
BOTTOM = 70
CONTENT_W = PAGE_W - 2 * MARGIN_X
FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"
COLORS = {
    "ink": (24, 34, 42),
    "blue": (20, 62, 96),
    "muted": (92, 105, 115),
    "rule": (205, 214, 221),
    "fill": (239, 243, 246),
}


def font(size: int):
    return ImageFont.truetype(FONT_PATH, size=size)


FONTS = {
    "title": font(42),
    "h2": font(31),
    "h3": font(25),
    "body": font(21),
    "table": font(17),
    "small": font(15),
}


def clean(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def wrap(draw: ImageDraw.ImageDraw, text: str, selected_font, width: int) -> list[str]:
    text = clean(text)
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textlength(candidate, font=selected_font) <= width or not current:
            current = candidate
        else:
            lines.append(current.rstrip())
            current = char.lstrip()
    if current:
        lines.append(current.rstrip())
    return lines


class Paginator:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for old in self.output_dir.glob("page-*.png"):
            old.unlink()
        self.pages: list[Path] = []
        self.page_no = 0
        self.image = None
        self.draw = None
        self.y = 0
        self.new_page()

    def new_page(self):
        if self.image is not None:
            self.save_page()
        self.page_no += 1
        self.image = Image.new("RGB", (PAGE_W, PAGE_H), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.draw.text((MARGIN_X, 30), "2026-07-17 基金午间研究报告", font=FONTS["small"], fill=COLORS["muted"])
        self.draw.line((MARGIN_X, 54, PAGE_W - MARGIN_X, 54), fill=COLORS["rule"], width=1)
        self.y = TOP

    def save_page(self):
        assert self.image is not None and self.draw is not None
        footer = str(self.page_no)
        self.draw.text((PAGE_W - MARGIN_X - self.draw.textlength(footer, font=FONTS["small"]), PAGE_H - 43), footer, font=FONTS["small"], fill=COLORS["muted"])
        path = self.output_dir / f"page-{self.page_no:02d}.png"
        self.image.save(path, optimize=True)
        self.pages.append(path)

    def ensure(self, height: int):
        if self.y + height > PAGE_H - BOTTOM:
            self.new_page()

    def add_block(self, text: str, kind: str = "body", prefix: str = ""):
        assert self.draw is not None
        selected = FONTS[kind]
        if kind == "title":
            color, line_h, before, after, indent = COLORS["blue"], 56, 10, 20, 0
        elif kind == "h2":
            color, line_h, before, after, indent = COLORS["blue"], 43, 18, 9, 0
        elif kind == "h3":
            color, line_h, before, after, indent = COLORS["blue"], 36, 13, 7, 0
        elif kind == "table":
            color, line_h, before, after, indent = COLORS["ink"], 25, 3, 3, 12
        else:
            color, line_h, before, after, indent = COLORS["ink"], 32, 4, 5, 0
        available = CONTENT_W - indent
        lines = wrap(self.draw, prefix + text, selected, available)
        total = before + len(lines) * line_h + after
        self.ensure(total)
        self.y += before
        for line in lines:
            self.draw.text((MARGIN_X + indent, self.y), line, font=selected, fill=color)
            self.y += line_h
        self.y += after

    def finish(self):
        self.save_page()
        return self.pages


def render_markdown(source: Path, output_dir: Path) -> list[Path]:
    pager = Paginator(output_dir)
    table_header: list[str] | None = None
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line == "<!-- PAGEBREAK -->":
            pager.new_page()
            table_header = None
            continue
        if line.startswith("# "):
            pager.add_block(line[2:], "title")
            continue
        if line.startswith("## "):
            pager.add_block(line[3:], "h2")
            table_header = None
            continue
        if line.startswith("### "):
            pager.add_block(line[4:], "h3")
            continue
        if line.startswith("| "):
            values = [clean(v.strip()) for v in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-+:?", value.replace(" ", "")) for value in values):
                continue
            if table_header is None:
                table_header = values
                pager.add_block("｜".join(values), "table", "表头｜")
            else:
                paired = "；".join(f"{table_header[i]}：{value}" for i, value in enumerate(values) if i < len(table_header))
                pager.add_block(paired, "table", "• ")
            continue
        table_header = None
        if line.startswith("- "):
            pager.add_block(line[2:], "body", "• ")
        elif re.match(r"^\d+\.\s", line):
            marker = line.split(".", 1)[0] + ". "
            pager.add_block(re.sub(r"^\d+\.\s", "", line), "body", marker)
        else:
            pager.add_block(line, "body")
    return pager.finish()


def build_docx(pages: list[Path], output: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    for idx, page in enumerate(pages):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1
        p.add_run().add_picture(str(page), width=Inches(8.5), height=Inches(11))
        if idx < len(pages) - 1:
            p.add_run().add_break(WD_BREAK.PAGE)
    props = doc.core_properties
    props.title = "2026-07-17 基金午间研究报告"
    props.subject = "中文可读图像页DOCX；可编辑源见同名Markdown"
    props.author = "Codex"
    doc.save(output)
    print(f"created={output} image_pages={len(pages)}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        raise SystemExit("usage: build_readable_image_docx.py input.md pages_dir output.docx")
    pages = render_markdown(Path(sys.argv[1]), Path(sys.argv[2]))
    build_docx(pages, Path(sys.argv[3]))
