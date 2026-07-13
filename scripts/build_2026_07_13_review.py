from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "data" / "holding-intake.csv"
TECH_INPUT = ROOT / "reports" / "2026-07-13-holding-technical-indicators.csv"
MODELED_HOLDINGS_INPUT = ROOT / "data" / "holdings.csv"
MODELED_ACTIONS_INPUT = ROOT / "reports" / "2026-07-13-modeled-redemption-plan.csv"
ASSUMPTIONS_INPUT = ROOT / "config" / "model-assumptions.json"
OUTPUT = ROOT / "reports" / "2026-07-13-基金晚间决策型复盘-增强版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "256D4A"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color="D0D5DD", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def set_font(run, size=11, bold=False, color=INK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, before=0, after=6, line=1.10, align=None) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_rich_paragraph(doc, parts, before=0, after=6, line=1.10, align=None):
    p = doc.add_paragraph()
    format_paragraph(p, before, after, line, align)
    for text, options in parts:
        run = p.add_run(text)
        set_font(run, **options)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_callout(doc, label: str, text: str, tone="blue") -> None:
    fills = {"blue": PALE_BLUE, "red": "FCE8E6", "gold": "FFF4CE", "green": "E6F4EA"}
    colors = {"blue": DARK_BLUE, "red": RED, "gold": GOLD, "green": GREEN}
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, fills[tone], "2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    format_paragraph(p, 1, 1, 1.10)
    r = p.add_run(f"{label}  ")
    set_font(r, 11, True, colors[tone])
    r = p.add_run(text)
    set_font(r, 11, False, INK)
    spacer = doc.add_paragraph()
    format_paragraph(spacer, 0, 2, 1.0)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.extend([rfonts, color, underline, size])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("第 ")
    set_font(r, 9, False, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.extend([begin, instr, separate, end])
    r2 = paragraph.add_run(" 页")
    set_font(r2, 9, False, MUTED)


def build_document() -> None:
    rows = list(csv.DictReader(INPUT.open(encoding="utf-8")))
    tech_rows = list(csv.DictReader(TECH_INPUT.open(encoding="utf-8")))
    tech_by_code = {r["fund_code"]: r for r in tech_rows}
    modeled_holdings = list(csv.DictReader(MODELED_HOLDINGS_INPUT.open(encoding="utf-8")))
    modeled_actions = list(csv.DictReader(MODELED_ACTIONS_INPUT.open(encoding="utf-8")))
    assumptions = json.loads(ASSUMPTIONS_INPUT.read_text(encoding="utf-8"))
    modeled_by_code = {r["fund_code"]: r for r in modeled_holdings}
    modeled_total = sum(float(r["market_value"]) for r in modeled_holdings)
    modeled_action_total = sum(float(r["modeled_market_value"]) for r in modeled_actions)
    modeled_action_fees = sum(float(r["modeled_redemption_fee"]) for r in modeled_actions)
    modeled_action_net = sum(float(r["modeled_net_proceeds"]) for r in modeled_actions)
    active = [r for r in rows if not r["status"].startswith("sold_")]
    active_total = sum(float(r["market_value"]) for r in active)
    active_cost = sum(float(r["estimated_cost"]) for r in active)
    active_gain = sum(float(r["holding_gain"]) for r in active)
    all_total = sum(float(r["market_value"]) for r in rows)
    qdii_total = sum(float(r["market_value"]) for r in active if "qdii" in r["cluster"])
    tech_total = sum(
        float(r["market_value"])
        for r in active
        if any(x in r["cluster"] for x in ("semiconductor", "tech", "communication", "cloud"))
    )

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    format_paragraph(header, 0, 0, 1.0)
    r = header.add_run("基金账户复盘  ·  研究记录")
    set_font(r, 9, True, MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    p = doc.add_paragraph()
    format_paragraph(p, 8, 4, 1.0)
    r = p.add_run("基金晚间决策型复盘")
    set_font(r, 24, True, "000000")
    p = doc.add_paragraph()
    format_paragraph(p, 0, 16, 1.0)
    r = p.add_run("2026年7月13日收盘后｜仓位底稿截至7月12日截图口径")
    set_font(r, 12, False, MUTED)

    metadata = [
        ("账户定位", "中国大陆场外公募基金；人工确认交易"),
        ("报告范围", "持仓核对、全市场信息面、主力资金、技术面、预测、教学与风险审计"),
        ("数据时点", "A股/港股截至7月13日收盘；美股仅截至7月10日收盘"),
        ("建模口径", "申购日统一按7月3日；份额由截图市值反推；满7日赎回费按0.5%"),
    ]
    for label, value in metadata:
        add_rich_paragraph(doc, [(f"{label}：", {"size": 10.5, "bold": True, "color": "000000"}), (value, {"size": 10.5, "bold": False, "color": "000000"})], after=2)

    add_heading(doc, "一、先复述你的仓位", 1)
    add_callout(
        doc,
        "迁移结论",
        f"底稿共19条。剔除“创新药已卖出、待确认”后，当前暂按18只在持，市值 {active_total:,.2f} 元，估算成本 {active_cost:,.2f} 元，截图口径浮亏 {abs(active_gain):,.2f} 元（{active_gain/active_cost:.2%}）。若把待确认卖出的1,421.58元也计入，截图总额为 {all_total:,.2f} 元。",
        "blue",
    )
    add_callout(doc, "7月13日模型更新", f"按7月12日截图市值除以当时最新正式净值反推份额，再乘以7月13日最新正式净值，18只在持基金的模型市值为 {modeled_total:,.2f} 元。若全部赎回并统一按0.5%估算，费用约 {modeled_total*0.005:,.2f} 元，净到账约 {modeled_total*0.995:,.2f} 元。", "green")
    add_rich_paragraph(doc, [("口头版复述：", {"size": 11, "bold": True, "color": DARK_BLUE}), ("你以沪深300作为最大单一底仓，同时叠加半导体材料设备、电子信息、通信设备、云计算和多只科技主动基金；海外部分由全球成长、全球产业升级、高端制造、纳指100及若干科技QDII组成。仓位数量多、科技因子重叠明显。", {"size": 11, "bold": False, "color": INK})], after=8)

    holdings = doc.add_table(rows=1, cols=6)
    headers = ["代码", "基金", "市值（元）", "占在持", "持有收益", "状态"]
    widths = [950, 3860, 1300, 1050, 1100, 1100]
    for i, text in enumerate(headers):
        cell = holdings.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        run = p.add_run(text)
        set_font(run, 9.2, True, "000000")
    set_repeat_table_header(holdings.rows[0])
    for rdata in rows:
        sold = rdata["status"].startswith("sold_")
        vals = [
            rdata["fund_code"],
            rdata["fund_name"],
            f"{float(rdata['market_value']):,.2f}",
            "—" if sold else f"{float(rdata['market_value'])/active_total:.2%}",
            f"{float(rdata['holding_gain']):+,.2f} / {float(rdata['holding_return']):+.2%}",
            "已卖待确认" if sold else "暂按在持",
        ]
        row = holdings.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(vals):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            run = p.add_run(text)
            color = MUTED if sold else (RED if i == 4 and float(rdata["holding_gain"]) < 0 else INK)
            set_font(run, 8.7, False, color)
    set_table_geometry(holdings, widths, 120)
    set_table_borders(holdings)

    add_heading(doc, "二、仓位结构与迁移质量", 1)
    structure = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["检查项", "当前口径", "项目约束", "结论"]):
        set_cell_shading(structure.rows[0].cells[i], PALE_BLUE)
        p = structure.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.5, True, "000000")
    checks = [
        ("持仓数量", "18只在持 + 1只待确认卖出", "最多3只", "明显不一致"),
        ("最大单只", "沪深300 25.30%", "不高于25%", "轻微超限"),
        ("显性科技主题", f"至少 {tech_total/active_total:.2%}", "主题不高于40%", "已超；实际或更高"),
        ("QDII", f"{qdii_total:,.2f}元 / {qdii_total/active_total:.2%}", "不高于25%", "口径内未超"),
        ("资金规模", f"在持市值 {active_total:,.2f}元", "配置本金30,000元", "配置可能已过期"),
        ("现金与回撤", "现金未知；无组合历史", "现金≥25%；回撤8%降仓", "无法核验"),
    ]
    for values in checks:
        row = structure.add_row()
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            color = RED if i == 3 and text not in ("口径内未超",) else INK
            set_font(p.add_run(text), 9.2, i == 0, color)
    set_table_geometry(structure, [1700, 3300, 1900, 2460], 120)
    set_table_borders(structure)
    add_callout(doc, "建模假设已落地", "所有在持基金统一按2026年7月3日申购、持有10天处理；份额由截图市值与7月12日前最新正式净值反推；赎回费采用用户指定的宽口径：未满7天1.5%，满7天0.5%。该口径足以生成金额草案，但不代表每只产品合同或销售平台的最终收费。现金余额与账户历史高点仍未知，只影响账户层面回撤，不阻断单只基金金额测算。", "gold")

    add_heading(doc, "三、今天的信息面复盘", 1)
    radar = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["信息维度", "当日事实", "市场含义", "持仓映射"]):
        set_cell_shading(radar.rows[0].cells[i], PALE_BLUE)
        p = radar.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.2, True, "000000")
    set_repeat_table_header(radar.rows[0])
    radar_rows = [
        ("A股全景", "三大指数下跌；超4600股下跌；成交2.82万亿元", "广泛风险收缩，不是少数权重拖累", "全部境内权益基金"),
        ("港股", "恒指+0.16%，恒科-0.96%；成交3095亿港元", "大盘防守强于科技，内部仍是风格切换", "沪深港云计算及港股暴露"),
        ("全球定价", "油价与美债收益率受地缘风险推升", "通胀与折现率压力不利高估值成长", "纳指、全球成长、半导体"),
        ("ETF中期资金", "7月6日至10日股票ETF净流入约828亿元", "此前有逆势承接，但尚未证明7月13日止跌", "半导体材料设备与宽基"),
        ("国内数据日程", "7月14日9:30有国民经济运行数据发布", "宏观预期差可能改变价值/成长相对强弱", "沪深300与主动成长"),
        ("盈利验证", "中报预告进入密集期，高位科技波动放大", "行情由题材叙事转向订单、利润与估值匹配", "电子、通信、集成电路"),
    ]
    for values in radar_rows:
        row = radar.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 8.8, i == 0, INK)
    set_table_geometry(radar, [1400, 2750, 3000, 2210], 120)
    set_table_borders(radar)
    add_heading(doc, "1. 外部冲击：地缘风险抬升油价与利率压力", 2)
    add_rich_paragraph(doc, [("事实。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("周末中东冲突再度升级，7月13日全球风险资产承压。AP报道亚洲交易时段韩国KOSPI下跌8.9%；美国10年期国债收益率升至4.59%，高于上周五的4.56%。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("因果链。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("冲突升级 → 原油与通胀风险上升 → 降息预期受压、长端利率抬升 → 高估值成长资产折现率上升。你的半导体、通信、电子信息和全球成长仓位对这一链条更敏感。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("反方解释。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("今天并非单纯由地缘事件造成。前期AI硬件拥挤、估值偏高和中报盈利验证压力，可能才是内部主因；外部冲击更像触发器和放大器。", {"size": 11, "bold": False, "color": INK})])

    add_heading(doc, "2. A股内部：科技拥挤交易进入盈利验证", 2)
    add_rich_paragraph(doc, [("事实。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("上周股票ETF净流入约828亿元，半导体材料设备相关ETF是主要吸金方向；但今天存储、PCB、先进封装、光纤光通信等高位方向集中回撤。7月又进入中报业绩预告密集期，资金从“叙事定价”转向订单、利润和估值匹配。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("市场验证。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("银行、石油天然气、煤炭、燃气等防御与资源方向相对占优，而电子元器件、半导体、电脑硬件和光通信明显承压，符合“高拥挤成长降温、资金向防御再平衡”的解释。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("明日事件。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("国家统计局日程显示7月14日9:30将发布国民经济运行相关数据。若数据或政策预期明显偏离市场预期，可能改变宽基与成长风格的相对表现。", {"size": 11, "bold": False, "color": INK})])

    add_heading(doc, "四、主力资金动向：不是简单的科技内轮动", 1)
    add_callout(doc, "先给结论", "按证券时报·数据宝的大单统计口径，7月13日更接近‘全市场大额资金普遍降风险、科技成为主要兑现区’，而不是资金从科技完整迁往另一个进攻板块。只有石油石化和农林牧渔两个申万一级行业为净流入，规模远小于全市场净流出；银行和中药虽然价格相对走强，但不能据此直接写成承接了全部流出资金。", "red")
    flow = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["观察口径", "7月13日读数", "判断", "下一步验证"]):
        set_cell_shading(flow.rows[0].cells[i], LIGHT)
        p = flow.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.3, True, "000000")
    set_repeat_table_header(flow.rows[0])
    flow_rows = [
        ("沪深两市主力资金", "净流出1147.08亿元", "大单口径出现显著风险收缩", "连续2—3日是否仍大额净流出"),
        ("行业广度", "31个行业中仅2个净流入、29个净流出", "接近集体撤退，不是健康的多板块轮动", "净流入行业能否扩至8个以上"),
        ("电子", "净流出520.93亿元", "是最主要兑现区，占全市场净流出约45%", "电子能否缩量止跌并转为净流入"),
        ("机械设备", "净流出112.59亿元", "科技制造抛压向设备扩散", "设备龙头相对沪深300是否止弱"),
        ("通信/电力设备/军工/计算机", "各净流出均超过60亿元", "成长风险并非局限于单一细分", "板块间是否出现持续性强弱分化"),
        ("逆势流入", "石油石化+2.62亿元；农林牧渔+0.96亿元", "有防御去向但承接规模很小", "后续是否连续流入且放量上涨"),
        ("尾盘", "主力净流出126.03亿元", "收盘前仍未出现全面回补", "次日尾盘是否转为净流入"),
    ]
    for values in flow_rows:
        row = flow.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 8.9, i == 0, RED if i == 1 and "净流出" in text else INK)
    set_table_geometry(flow, [1800, 2200, 2900, 2460], 120)
    set_table_borders(flow)
    add_rich_paragraph(doc, [("对持仓的含义。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("你的电子、通信、半导体、集成电路、云计算和多只科技主动基金暴露重叠，因此电子与成长板块的大额净流出会被组合重复放大。沪深300、银行权重和极小的黄金仓位能提供的对冲有限。单日资金流不构成卖出理由，但若连续2—3日出现行业广度恶化、电子继续大额净流出且相对沪深300走弱，‘只是一次获利回吐’的解释将明显变弱。", {"size": 11, "bold": False, "color": INK})], before=8)
    add_rich_paragraph(doc, [("口径警告。", {"size": 10, "bold": True, "color": MUTED}), ("这里的‘主力资金’通常是行情软件按主动买卖方向和大单/超大单阈值推算的成交统计，不是交易所披露的机构真实账户流水。不同平台阈值、样本和算法不同，所以看方向、行业广度和持续性，比迷信某一个亿元数字更可靠。", {"size": 10, "bold": False, "color": MUTED})], after=8)

    add_heading(doc, "五、今天的技术面复盘", 1)
    market = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["观察对象", "7月13日表现", "技术含义", "对应持仓"]):
        set_cell_shading(market.rows[0].cells[i], LIGHT)
        p = market.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.5, True, "000000")
    observations = [
        ("沪深300", "-1.79%", "跌幅小于成长指数，权重价值相对抗跌", "沪深300联接"),
        ("上证指数", "-2.06%，3913.79", "早盘-0.75%后继续走弱，日内承接不足", "宽基/主动A股"),
        ("深证成指", "-3.48%，14522.85", "成长与科技权重拖累明显", "电子、半导体、通信"),
        ("创业板指", "-3.10%，3723.52", "高弹性成长弱于沪深300约1.31个百分点", "主动成长"),
        ("科创综指", "-4.36%", "高估值硬科技是主要风险释放区", "半导体材料设备"),
        ("市场宽度", "超4600股下跌；逾170股跌停", "不是少数权重拖累，而是广泛风险收缩", "几乎全部A股仓位"),
        ("两市成交", "2.8178万亿元；缩量5708亿元", "跌价缩量说明抛压与接盘同时退潮，尚非放量恐慌出清", "观察次日修复质量"),
        ("恒生指数", "+0.16%；成交3095亿港元", "大盘横住，但内部风格分化", "沪港深与港股QDII"),
        ("恒生科技", "-0.96%", "科技弱于恒指1.12个百分点", "云计算、全球成长"),
        ("美股上周五", "纳指+0.3%，周+1.7%", "海外趋势截至周五仍偏强；周一未收盘", "纳指及海外科技QDII"),
    ]
    for values in observations:
        row = market.add_row()
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            color = RED if i == 1 and text.startswith("-") else INK
            set_font(p.add_run(text), 9.1, i == 0, color)
    set_table_geometry(market, [1500, 2200, 3500, 2160], 120)
    set_table_borders(market)

    add_rich_paragraph(doc, [("技术结论。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("今天是“低开—弱反弹—继续走低”的广泛风险收缩日，且科技成长显著弱于宽基。缩量大跌不等于止跌：它说明买盘退缩，但尚未出现足以确认筹码充分交换的放量企稳。短线需要等待价格、成交和市场宽度同时改善。", {"size": 11, "bold": False, "color": INK})], before=8)
    add_heading(doc, "持仓基金的净值技术指标", 2)
    add_callout(doc, "数据已主动补齐", "已从公开基金净值页面同步19只基金共23,428条正式历史净值。12只境内基金最新净值日期为7月13日；7只QDII最新为7月10日，属于跨时区净值确认时滞，不当作7月13日盘中价格。", "green")
    add_rich_paragraph(doc, [("组合读数。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("剔除已卖待确认的创新药后，18只在持基金中只有3只站上20日均线，对应截图市值7,484.25元、占在持市值18.95%；其余15只占81.05%，均在20日均线下。RSI集中在40.12—59.10，没有极端超买或超卖，说明当前更像趋势降温与结构分化，而不是单一指标给出的反转点。", {"size": 11, "bold": False, "color": INK})], after=8)

    tech = doc.add_table(rows=1, cols=7)
    for i, text in enumerate(["代码", "基金（简称）", "最新净值/日期", "相对MA20", "相对MA60", "RSI14", "结构"]):
        set_cell_shading(tech.rows[0].cells[i], LIGHT)
        p = tech.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 8.8, True, "000000")
    set_repeat_table_header(tech.rows[0])
    short_names = {
        "007339": "沪深300联接C", "021894": "半导体材料设备C", "010004": "电子信息产业C",
        "012922": "全球成长QDII C", "021528": "财通成长优选C", "007818": "通信设备联接C",
        "019024": "信息行业精选C", "024481": "财通品质甄选C", "017731": "全球产业升级QDII C",
        "014565": "创新药50联接C", "019331": "沪深港云计算C", "016665": "全球高端制造QDII C",
        "006479": "纳指100 QDII C", "016371": "信澳业绩驱动C", "006503": "集成电路产业C",
        "002891": "移动互联QDII", "539002": "新兴市场QDII A", "006373": "全球科技互联QDII A",
        "021873": "沪深港黄金产业A",
    }
    for source_row in rows:
        t = tech_by_code[source_row["fund_code"]]
        sold = source_row["status"].startswith("sold_")
        vals = [
            source_row["fund_code"], short_names[source_row["fund_code"]],
            f"{float(t['latest_nav']):.4f}\n{t['latest_date']}",
            f"{float(t['nav_vs_ma20']):+.2%}", f"{float(t['nav_vs_ma60']):+.2%}",
            t["rsi14"], ("已卖待确认；" if sold else "") + t["technical_state"],
        ]
        row = tech.add_row()
        for i, text in enumerate(vals):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (1, 6) else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            value_for_color = float(t["nav_vs_ma20"]) if i == 3 else None
            color = MUTED if sold else (RED if value_for_color is not None and value_for_color < 0 else INK)
            set_font(p.add_run(text), 8.0, i == 0, color)
    set_table_geometry(tech, [820, 1900, 1420, 1100, 1100, 900, 2120], 120)
    set_table_borders(tech)

    add_heading(doc, "收益、回撤与相对强弱", 2)
    performance = doc.add_table(rows=1, cols=7)
    for i, text in enumerate(["代码", "20日收益", "60日收益", "60日最大回撤", "历史最大回撤", "相对沪深300 20日", "相对沪深300 60日"]):
        set_cell_shading(performance.rows[0].cells[i], PALE_BLUE)
        p = performance.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 8.2, True, "000000")
    set_repeat_table_header(performance.rows[0])
    for source_row in rows:
        t = tech_by_code[source_row["fund_code"]]
        sold = source_row["status"].startswith("sold_")
        vals = [
            source_row["fund_code"], f"{float(t['return_20obs']):+.2%}", f"{float(t['return_60obs']):+.2%}",
            f"-{float(t['drawdown_60obs']):.2%}", f"-{float(t['max_drawdown_full_history']):.2%}",
            f"{float(t['relative_20obs_vs_007339']):+.2%}", f"{float(t['relative_60obs_vs_007339']):+.2%}",
        ]
        row = performance.add_row()
        for i, text in enumerate(vals):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            color = MUTED if sold else (RED if i in (1, 2, 5, 6) and float(text.replace("%", "")) < 0 else INK)
            set_font(p.add_run(text), 8.2, i == 0, color)
    set_table_geometry(performance, [850, 1200, 1200, 1500, 1500, 1555, 1555], 120)
    set_table_borders(performance)
    add_rich_paragraph(doc, [("口径说明。", {"size": 9.5, "bold": True, "color": MUTED}), ("MA20/MA60为最近20/60个已公布净值观察值的简单均线；RSI14使用Wilder平滑；回撤由基金净值序列计算，不等同于你的账户回撤；相对强弱为基金同期收益减去007339沪深300联接C同期收益。QDII日期按其已公布净值，不用7月13日海外盘中行情回填。", {"size": 9.5, "bold": False, "color": MUTED})], before=6, after=8, line=1.05)
    add_callout(doc, "持仓技术结论", "半导体材料设备联接C仍是最强持仓：净值高于MA20约1.48%、高于MA60约29.41%，RSI 59.10，20日相对沪深300强约29.66个百分点。云计算联接C和纳指100联接C也在MA20上方。相对较弱的在持基金是集成电路产业C、财通品质甄选C和财通成长优选C，分别低于MA20约17.73%、16.14%和15.91%。这些是趋势状态，不是脱离费用与持有期的即时买卖指令。", "gold")

    add_heading(doc, "六、预测：概率、验证与失效条件", 1)
    add_callout(doc, "预测结论", "下一交易日的基准判断是‘先尝试技术修复、随后继续分化’，不是确认见底；未来1—2周的基准判断是‘高位科技去拥挤与盈利验证并行，指数震荡、个股分化加大’。这是基于7月13日收盘信息作出的条件概率，后续资金流、价格和成交会动态改写概率。", "blue")
    add_heading(doc, "下一交易日情景", 2)
    next_day = doc.add_table(rows=1, cols=5)
    for i, text in enumerate(["情景", "概率", "因果链", "验证信号", "持仓动作"]):
        set_cell_shading(next_day.rows[0].cells[i], PALE_BLUE)
        p = next_day.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.0, True, "000000")
    set_repeat_table_header(next_day.rows[0])
    next_day_rows = [
        ("基准：弱修复后分化", "50%", "普跌后的技术反抽 + 资金仍谨慎", "沪指守3900；上涨家数改善但不足压倒性；电子净流出明显收窄", "不追反弹；观察最强持仓能否率先转强"),
        ("偏弱：继续释放", "30%", "尾盘未回补 + 科技拥挤继续松动", "沪指有效跌破3900；不足1000股上涨；电子/计算机继续放量净流出", "优先压降重复科技暴露，执行前核费用与到账"),
        ("偏强：放量修复", "20%", "宏观预期改善 + ETF/主动资金承接", "上涨家数超过3500；电子与通信转净流入；指数接近日内高位收盘", "只讨论分批调整，不在单日急涨中追高"),
    ]
    for values in next_day_rows:
        row = next_day.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 8.5, i == 0, INK)
    set_table_geometry(next_day, [1600, 800, 2300, 2800, 1860], 120)
    set_table_borders(next_day)
    add_heading(doc, "未来1—2周情景", 2)
    scenarios = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["情景", "主观概率", "验证信号", "当前动作"]):
        set_cell_shading(scenarios.rows[0].cells[i], PALE_BLUE)
        p = scenarios.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 9.5, True, "000000")
    data = [
        ("基准：科技去拥挤、指数震荡", "50%", "资金流出收窄但科技内部强弱分化；中报决定持续性", "保留强趋势仓位，审查弱趋势与重复暴露"),
        ("偏强：资金回流且盈利兑现", "25%", "科技连续3日强于沪深300，成交与上涨家数同步改善", "仅在验证后分批调整，不提前押注V形反转"),
        ("偏弱：外部利率与内部盈利共振", "25%", "油价/美债再升，核心公司盈利预期下修，科技放量破位", "优先降低主题集中度并执行8%回撤纪律"),
    ]
    for values in data:
        row = scenarios.add_row()
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 9.2, i == 0, INK)
    set_table_geometry(scenarios, [2550, 1200, 3250, 2360], 120)
    set_table_borders(scenarios)
    add_rich_paragraph(doc, [("最强反方观点。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("上周股票ETF仍大幅净流入，7月13日可能只是极端拥挤交易的一次快速清洗；若资金很快回流半导体材料设备、核心公司中报继续上修，当前谨慎预测会低估V形修复。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("推翻基准预测的证据。", {"size": 11, "bold": True, "color": DARK_BLUE}), ("偏强方向：连续3个交易日电子/通信主力净流入、上涨家数占优、科技相对沪深300转强。偏弱方向：沪指跌破关键位后不能收回，成交放大、行业净流出广度维持高位且盈利预期下调。任何一组证据成立，都要重新分配情景概率。", {"size": 11, "bold": False, "color": INK})])
    add_heading(doc, "七、行动草案与风险否决", 1)
    add_heading(doc, "按统一日期与费率生成的赎回草案", 2)
    add_callout(doc, "草案汇总", f"确定性规则触发14只退出审查，模型市值合计 {modeled_action_total:,.2f} 元；按0.5%估算赎回费用 {modeled_action_fees:,.2f} 元，净到账约 {modeled_action_net:,.2f} 元。未触发的4只合计约 {modeled_total-modeled_action_total:,.2f} 元。金额均为研究草案，不会自动下单。", "red")
    action_table = doc.add_table(rows=1, cols=7)
    for i, text in enumerate(["代码", "基金（简称）", "模型市值", "费率", "估算费用", "净到账", "触发依据"]):
        set_cell_shading(action_table.rows[0].cells[i], "FCE8E6")
        p = action_table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, 0, 0, 1.0)
        set_font(p.add_run(text), 8.5, True, "000000")
    set_repeat_table_header(action_table.rows[0])
    for action in modeled_actions:
        evidence = action["evidence"].replace("；", "\n")
        vals = [
            action["fund_code"], short_names[action["fund_code"]],
            f"{float(action['modeled_market_value']):,.2f}", f"{float(action['modeled_redemption_fee_rate']):.2%}",
            f"{float(action['modeled_redemption_fee']):,.2f}", f"{float(action['modeled_net_proceeds']):,.2f}", evidence,
        ]
        row = action_table.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(vals):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (1, 6) else WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 8.0, i == 0, RED if i == 6 else INK)
    set_table_geometry(action_table, [800, 1900, 1150, 800, 1050, 1150, 2510], 120)
    set_table_borders(action_table)
    add_callout(doc, "风险官结论：已生成金额草案", "不再以申购日期、估算份额或宽口径费率缺失为由停止计算。当前草案使用7月3日统一申购日期和0.5%赎回费；真正提交前只需知道平台最终费率可能与模型不同，并由你本人确认。", "green")
    add_rich_paragraph(doc, [("推翻“继续震荡”判断的证据：", {"size": 11, "bold": True, "color": DARK_BLUE}), ("连续交易日出现科技相对沪深300转强、成交额回升、上涨家数显著占优，并由中报订单与利润兑现配合。", {"size": 11, "bold": False, "color": INK})])
    add_rich_paragraph(doc, [("推翻“只是技术性调整”的证据：", {"size": 11, "bold": True, "color": DARK_BLUE}), ("科技主题持续放量下跌、核心持仓盈利预期下修、AI资本开支或半导体设备订单出现实质性下调，同时账户真实回撤触发8%纪律线。", {"size": 11, "bold": False, "color": INK})])

    add_heading(doc, "八、今日知识点：怎样正确使用‘主力资金’", 1)
    add_callout(doc, "一句话教学", "主力资金不是机构账户流水，而是按成交方向与订单大小推算的市场行为指标。它适合回答‘风险偏好正在扩散还是收缩’，不适合单独回答‘明天一定涨还是跌’。", "gold")
    teaching = [
        ("第一层：价格", "资金流出时价格是否同步下跌？7月13日电子大额净流出且板块下跌，信号相互确认。"),
        ("第二层：成交", "下跌是放量恐慌还是缩量退潮？当天总成交缩量，说明买卖双方都在退缩，不能直接视为充分出清。"),
        ("第三层：广度", "只有电子下跌，还是多数行业一起流出？当天29个申万一级行业净流出，风险收缩具有广度。"),
        ("第四层：持续性", "单日异常容易反转；连续2—3日同方向，且相对强弱不修复，可信度才明显上升。"),
        ("放到你的持仓", "半导体材料设备C仍在MA20上方，说明它虽受板块冲击但趋势相对更强；集成电路产业C远低于MA20，若资金流继续恶化，应优先审查弱趋势和重复暴露，而不是把所有科技基金一刀切。"),
    ]
    for label, value in teaching:
        add_rich_paragraph(doc, [(f"{label}。", {"size": 11, "bold": True, "color": DARK_BLUE}), (value, {"size": 11, "bold": False, "color": INK})], after=5)
    add_callout(doc, "可复用检查句", "以后看到‘主力净流出’先问四个问题：价格确认了吗？成交放大了吗？行业广度有多大？持续了几天？四项越一致，信号越可靠。", "green")

    add_heading(doc, "九、下一步需要补齐的信息", 1)
    required = [
        ("每只基金", "模型已反推份额与成本净值；如平台份额不同，之后用实值覆盖模型值"),
        ("交易状态", "创新药赎回申请时间、确认份额/净值、到账金额与费用"),
        ("账户层面", "当前现金余额、账户历史高点、是否还有未迁移基金或在途交易"),
        ("费用与限制", "已按宽口径0.5%生成草案；平台最终费率仅作为执行前偏差提醒"),
    ]
    req_table = doc.add_table(rows=0, cols=2)
    for label, value in required:
        row = req_table.add_row()
        set_cell_shading(row.cells[0], LIGHT)
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, 0, 0, 1.0)
            set_font(p.add_run(text), 9.5, i == 0, INK)
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(req_table, [1900, 7460], 120)
    set_table_borders(req_table)

    doc.add_page_break()
    add_heading(doc, "十、来源与口径", 1)
    sources = [
        ("本地持仓底稿", "data/holding-intake.csv（截图迁移，时点2026-07-12）", None),
        ("建模假设", "config/model-assumptions.json（7月3日统一申购；满7日按0.5%赎回费）", None),
        ("基金净值历史", "天天基金公开净值页面；19只基金逐条URL与抓取时间已保存于data/nav-history.csv", "https://fund.eastmoney.com/007339.html"),
        ("A股收盘与成交", "每日经济新闻｜7月13日A股收盘", "https://www.nbd.com.cn/articles/2026-07-13/4469407.html"),
        ("宽基指数表现", "金融界｜7月13日宽基指数", "https://finance.jrj.com.cn/2026/07/13151657779197.shtml"),
        ("A股跌幅与原因", "东方财富｜A股大跌，发生了什么", "https://finance.eastmoney.com/a/202607133804134796.html"),
        ("ETF资金流", "东方财富｜股票ETF单周净流入近830亿元", "https://finance.eastmoney.com/a/202607133803694914.html"),
        ("主力资金全景", "证券时报·数据宝｜7月13日行业及尾盘主力资金", "https://finance.sina.com.cn/stock/zqgd/2026-07-13/doc-inihscme5765600.shtml"),
        ("主力资金方法", "东方财富资金流向数据中心及大单统计口径说明", "https://data.eastmoney.com/zjlx/"),
        ("两融与周度资金", "证券时报｜热点轮动加快，主力资金与两融资金现分歧", "https://www.stcn.com/article/detail/4012463.html"),
        ("港股收盘", "新浪财经/智通｜7月13日港股收盘", "https://finance.sina.com.cn/stock/hkstock/hkstocknews/2026-07-13/doc-inihscma8936694.shtml"),
        ("全球市场与利率", "AP｜7月13日油价、亚洲股市与美债", "https://apnews.com/article/2d6744b09c68b5473d0bc8584b89e60e"),
        ("美股上周五", "AP｜7月10日美国主要指数", "https://apnews.com/article/d3c5b8171e25e98da62831181de9a666"),
        ("统计数据日程", "国家统计局｜2026年主要统计信息发布日程", "https://www.stats.gov.cn/xw/tjxw/tzgg/202512/t20251224_1962137.html"),
    ]
    for label, text, url in sources:
        p = doc.add_paragraph()
        format_paragraph(p, 0, 4, 1.08)
        set_font(p.add_run(f"{label}："), 10, True, INK)
        if url:
            add_hyperlink(p, text, url)
        else:
            set_font(p.add_run(text), 10, False, INK)

    add_rich_paragraph(doc, [("免责声明：", {"size": 9.5, "bold": True, "color": MUTED}), ("本报告用于个人研究与仓位核对，不承诺收益，不替代基金销售平台的正式净值、费率与交易确认。场外基金按未知价申赎，QDII还存在时区与净值确认时滞；任何交易需由你本人核对并手工确认。", {"size": 9.5, "bold": False, "color": MUTED})], before=10, after=0, line=1.05)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
