from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "2026-07-15"
NOW = "2026-07-15T12:05:00+08:00"
DATA_CUTOFF = "2026-07-15T11:59:44+08:00"
DECISION_ID = "DEC-20260715-NOON-001"
DOCX = ROOT / "reports" / f"{REPORT_DATE}-基金午报.docx"
MD = ROOT / "reports" / f"{REPORT_DATE}-基金午报.md"


MARKET = {
    "sse": "3963.90（-0.08%）",
    "csi300": "4798.61（+0.04%）",
    "star50": "1936.06（-3.67%）",
    "chinext": "3837.23（-0.36%）",
    "csi500": "8205.12（-0.86%）",
    "csi1000": "7870.00（-0.70%）",
    "turnover": "沪深京约1.1228万亿元，线性外推约2.25万亿元，低于3万亿元强修复阈值",
    "hk": "恒生指数+1.47%，恒生科技+1.39%（11:59）",
    "us": "隔夜道指+0.02%，纳指+0.90%，标普500+0.38%",
}

ESTIMATES = {
    "021894": ("易方达半导体设备ETF联接C", "-5.23%"),
    "010004": ("景顺长城电子信息产业股票C", "-3.42%"),
    "006503": ("财通集成电路产业股票C", "-2.32%"),
    "007818": ("国泰中证全指通信设备ETF联接C", "-1.56%"),
    "019024": ("易方达信息行业精选股票C", "-0.80%"),
    "019331": ("华泰柏瑞中证沪港深云计算产业ETF联接C", "-0.19%"),
    "007339": ("易方达沪深300ETF联接C", "+0.05%"),
}

SOURCES = [
    "项目自检：node fund-investment/scripts/check-project-context.mjs --root fund-investment，2026-07-15运行，状态ready_with_warnings。",
    "正式净值：data/nav-history.csv，同步脚本sync-nav.mjs，抓取时间2026-07-15T03:53:29Z，来源为东方财富基金公开数据。",
    "盘中指数：腾讯行情 https://qt.gtimg.cn/q=sh000001,sh000300,sh000688,sz399006,sh000905,sh000852,sz399001,bj899050，抓取时间2026-07-15 11:59 CST。",
    "盘中估算净值：东方财富基金估值接口 https://fundgz.1234567.com.cn/ ，估算时间2026-07-15 11:30。",
    "港股与海外：腾讯行情 https://qt.gtimg.cn/q=r_hkHSI,r_hkHSCEI,r_hkHSTECH,usDJI,usIXIC,usINX，抓取时间2026-07-15 11:59。",
    "央行公开市场：人民银行公开市场业务交易公告[2026]第135号，https://www.pbc.gov.cn/zhengcehuobisi/125207/125213/125431/125475/2026071508544082234/index.html。",
    "WAIC日程：外交部外事日程 https://www.fmprc.gov.cn/wjdt_674879/wsrc_674883/202607/t20260713_11980587.shtml；上海市政府新闻 https://www.shanghai.gov.cn/nw4411/20260708/ba4c8e75f2744b43a6080ebb82a3aab2.html。",
    "统计发布日程：国家统计局2026年主要统计信息发布日程表 https://www.stats.gov.cn/sj/fbrc/bnxxfb/。",
    "长鑫科技发行信息：data/catalyst-ledger.csv既有记录及长鑫科技发行公告转载；发行价8.66元、7月16日申购、T+2缴款等只作为事件时钟，不作为基金成交价格。",
]


def read_csv(rel: str) -> list[dict[str, str]]:
    with (ROOT / rel).open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(rel: str, rows: list[dict[str, str]], fieldnames: list[str]) -> None:
    with (ROOT / rel).open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def latest_navs() -> dict[str, dict[str, str]]:
    latest: dict[str, dict[str, str]] = {}
    for row in read_csv("data/nav-history.csv"):
        if row.get("is_estimate") == "true":
            continue
        code = row["fund_code"]
        key = (row["date"], row["fetched_at"])
        if code not in latest or key > (latest[code]["date"], latest[code]["fetched_at"]):
            latest[code] = row
    return latest


def holdings_rows() -> tuple[list[dict[str, str]], float]:
    funds = {row["fund_code"]: row for row in read_csv("data/funds.csv")}
    intake = {row["fund_code"]: row for row in read_csv("data/holding-intake.csv")}
    navs = latest_navs()
    rows = []
    total = 0.0
    for h in read_csv("data/holdings.csv"):
        code = h["fund_code"]
        nav = navs.get(code, {})
        shares = float(h.get("shares") or 0)
        nav_value = float(nav.get("nav") or h.get("confirmed_nav") or 0)
        market_value = shares * nav_value
        total += market_value
        fund = funds.get(code, {})
        rows.append({
            "fund_code": code,
            "fund_name": fund.get("fund_name", code),
            "shares": f"{shares:.2f}",
            "nav_date": nav.get("date", ""),
            "nav": f"{nav_value:.4f}",
            "market_value": f"{market_value:.2f}",
            "cluster": intake.get(code, {}).get("cluster", fund.get("theme", "")),
            "is_qdii": fund.get("is_qdii", "false"),
        })
    return rows, total


def append_catalysts() -> None:
    path = "data/catalyst-ledger.csv"
    rows = read_csv(path)
    fieldnames = list(rows[0].keys())
    existing = {row["catalyst_id"] for row in rows}
    additions = []
    if "CAT-20260715-PBOC-OMO" not in existing:
        additions.append({
            "catalyst_id": "CAT-20260715-PBOC-OMO",
            "first_published_at": "2026-07-15T09:20:31+08:00",
            "event_start": "2026-07-15",
            "event_end": "2026-07-15",
            "title": "中国人民银行开展4265亿元7天期逆回购操作",
            "category": "monetary_liquidity/OMO",
            "verification_status": "official_observed",
            "confidence": "high",
            "novelty": "low",
            "expected_direction": "neutral_to_positive",
            "expected_horizon": "当日至3日",
            "beneficiary_chain": "公开市场流动性>短端资金预期>风险偏好与成交额>宽基和成长风格",
            "portfolio_mapping": "易方达沪深300ETF联接C;主动权益;半导体与通信主题",
            "priced_in_status": "partially_traded_intraday",
            "bull_case": "流动性投放与港股风险偏好配合，支撑宽基守住3900并缓和恐慌赎回",
            "bear_case": "逆回购是例行流动性工具，不能抵消半导体IPO兑现和科技拥挤出清",
            "validation_signals": "成交额维持2.6万亿元以上;宽基不破3900;电子和通信资金流出收窄",
            "invalidation_signals": "成交放大但科技继续高开低走;电子主力净流出超过300亿元;宽基跌破3900",
            "next_checkpoint": "2026-07-15 14:30与14:45复核成交额、宽基和电子/通信资金",
            "primary_source": "https://www.pbc.gov.cn/zhengcehuobisi/125207/125213/125431/125475/2026071508544082234/index.html",
            "secondary_source": "",
            "last_reviewed_at": NOW,
            "status": "watch",
        })
    if "CAT-20260715-CXMT-INTRADAY-VALIDATION" not in existing:
        additions.append({
            "catalyst_id": "CAT-20260715-CXMT-INTRADAY-VALIDATION",
            "first_published_at": "2026-07-15T11:30:00+08:00",
            "event_start": "2026-07-15",
            "event_end": "2026-07-16",
            "title": "长鑫发行价落地后半导体设备材料链午间明显跑输宽基",
            "category": "IPO/market_validation/semiconductor_rotation",
            "verification_status": "market_observed",
            "confidence": "medium",
            "novelty": "medium",
            "expected_direction": "negative_for_overlapping_semiconductor_exposure",
            "expected_horizon": "当日至3日",
            "beneficiary_chain": "长鑫发行市值锚>科技风险预算腾挪>存储/设备材料内部分化>半导体基金估算净值承压",
            "portfolio_mapping": "易方达半导体设备ETF联接C;景顺长城电子信息产业股票C;财通集成电路产业股票C;国泰中证全指通信设备ETF联接C",
            "priced_in_status": "being_traded_intraday",
            "bull_case": "若14:30后半导体设备材料收复相对沪深300劣势且通信/云计算未同步走弱，说明只是早盘兑现",
            "bear_case": "若半导体、通信、云计算同步弱于沪深300且上涨家数低于1000，虹吸从半导体扩散到科技主线",
            "validation_signals": "021894估算跌幅超过沪深300 5个百分点;半导体成分股多数下跌;科创50弱于沪深300",
            "invalidation_signals": "14:45半导体设备材料强于沪深300至少1个百分点;通信和云计算转强;资金流出转入",
            "next_checkpoint": "2026-07-15 14:30/14:45按长鑫行动卡复核",
            "primary_source": "https://fundgz.1234567.com.cn/",
            "secondary_source": "https://qt.gtimg.cn/",
            "last_reviewed_at": NOW,
            "status": "watch",
        })
    if additions:
        for item in additions:
            rows.append({key: item.get(key, "") for key in fieldnames})
        write_csv(path, rows, fieldnames)


def update_decision_journal() -> None:
    path = "data/decision-journal.csv"
    rows = read_csv(path)
    fieldnames = list(rows[0].keys())
    for row in rows:
        if row.get("evaluation_due_date") and row["evaluation_due_date"] <= REPORT_DATE and row.get("evaluation_status") == "pending":
            row["evaluation_status"] = "pending_formal_nav_or_confirmation"
            row["reviewed_at"] = NOW
            row["lesson"] = (
                "到期复核已检查：2026-07-14正式净值可得，但2026-07-15正式净值、平台确认、"
                "真实费用和到账仍未可得；暂不改参数。"
            )
            row["error_type"] = row.get("error_type") or "data_quality"
    if not any(row["decision_id"] == DECISION_ID for row in rows):
        new = {key: "" for key in fieldnames}
        new.update({
            "decision_id": DECISION_ID,
            "record_origin": "native_realtime",
            "report_time": NOW,
            "data_cutoff": DATA_CUTOFF,
            "decision_for_date": REPORT_DATE,
            "horizon_days": "3",
            "scope": "portfolio",
            "original_action": "wait now; conditional semiconductor-overlap redemption before cutoff only if 14:45 weakness persists",
            "target_funds": "021894;010004;006503;007818;019331;019024",
            "position_size": "0% now; 021894 33% on yellow/orange semiconductor confirmation; 021894 50% plus 007818 20% only if spread to communication; no subscription",
            "up_probability": "0.25",
            "flat_probability": "0.45",
            "down_probability": "0.30",
            "thesis": "SSE stays above 3900 and Hong Kong is positive, but STAR50 and semiconductor fund estimates show CXMT-related tech rotation; treat as structural risk, not account-wide liquidation",
            "trigger_conditions": "14:45 semiconductor-equipment/materials weaker than CSI300 by >2pp and large-order flow remains negative; upgrade if communication weaker by >1.5pp or advancers<1000",
            "invalidation_conditions": "SSE holds above 3900; semiconductor-equipment/materials recover to within 1pp of CSI300 or turn stronger; communication/cloud stabilize; platform fee/status conflicts",
            "execution_deadline": "2026-07-15T15:00:00+08:00",
            "expected_nav_date": "2026-07-15 if ordinary domestic OTC redemption is accepted before actual platform cutoff; QDII follows product calendar",
            "evaluation_due_date": "2026-07-18",
            "evaluation_status": "pending",
            "evidence_sources": "reports/2026-07-15-基金午报.docx;reports/2026-07-15-长鑫虹吸风险行动卡.md;data/catalyst-ledger.csv;data/nav-history.csv;https://qt.gtimg.cn/;https://fundgz.1234567.com.cn/;https://www.pbc.gov.cn/;https://www.fmprc.gov.cn/;https://www.shanghai.gov.cn/",
        })
        rows.append(new)
    write_csv(path, rows, fieldnames)


def make_markdown() -> str:
    holdings, total = holdings_rows()
    domestic = [r for r in holdings if r["is_qdii"] != "true"]
    qdii = [r for r in holdings if r["is_qdii"] == "true"]
    tech_mv = sum(float(r["market_value"]) for r in holdings if any(k in r["cluster"] for k in ["semiconductor", "communication", "cloud", "tech", "electronic", "integrated"]))
    qdii_mv = sum(float(r["market_value"]) for r in qdii)
    lines = []
    lines.append(f"# {REPORT_DATE} 基金午间研究报告")
    lines.append("")
    lines.append(f"- 生成时间：{NOW}")
    lines.append(f"- 数据截止：{DATA_CUTOFF}")
    lines.append("- 账户范围：只研究大陆销售平台可申赎的场外开放式基金；盘中指数、资金流和估算净值只用于决定是否在截止前提交申请，不是成交价格。")
    lines.append("- 项目自检：ready_with_warnings；组合历史为空，无法计算真实账户回撤；持仓市值与配置资金相差超过10%，现金比例不可直接沿用配置值。")
    lines.append("")
    lines.append("## 只看这一页：行动结论")
    lines.append("")
    lines.append("1. **今天15:00前：现在等待；14:45仍弱才提交结构性赎回，不申购。** 优先对象是易方达半导体设备ETF联接C：若14:45半导体设备材料仍弱于沪深300超过2个百分点、且资金流未收窄，提交约33%可赎回份额（模型约548.55份，按2026-07-14正式净值估算毛额约1617.29元，0.5%费率口径约8.09元）。若通信也弱于沪深300超过1.5个百分点或上涨家数低于1000，升级为易方达半导体设备ETF联接C 50%（约831.13份，毛额约2450.60元，0.5%费率约12.25元），国泰中证全指通信设备ETF联接C再审查20%（约96.05份，毛额约409.44元，费率待平台确认）。景顺长城电子信息产业股票C已有7月14日submitted赎回，不重复提交；只确认是否受理。")
    lines.append("2. **预计按哪天净值成交：** 普通境内场外基金若在平台实际截止前受理，通常按2026-07-15收盘后正式净值确认；15:00后通常顺延下一开放日。QDII不参与今天A股盘中动作，按各自产品日历滞后披露。")
    lines.append("3. **是否需要立即跑：** 不需要系统性撤退；这是科技内部结构性降仓预案。立即跑的证据必须是沪指跌破3900且14:45不能收复、上涨家数少于1000、电子/半导体资金扩散性净流出。")
    lines.append("4. **午后及下一交易日组合概率：** 收涨25%、震荡45%、收跌30%，合计100%。")
    lines.append("5. **今天只需要做的一件事：** 14:30初检、14:45终检易方达半导体设备ETF联接C是否触发33%或50%条件；平台费率/可赎份额不一致就等待下一开放日。")
    lines.append("6. **五个必须关注指标：** 上证3900；上涨家数3500/1000；半导体设备材料相对沪深300差值（强于+1个百分点取消、弱于-2个百分点触发）；通信相对沪深300（弱于-1.5个百分点升级）；成交额进度（14:45线性低于2.813万亿元则资金未确认强修复）。")
    lines.append("7. **推翻结论的证据：** 14:45半导体设备材料重新强于沪深300至少1个百分点，通信和云计算没有同步跑输，且电子/半导体资金流出收窄或转入。")
    lines.append("")
    lines.append("## 持仓复述")
    lines.append("")
    lines.append(f"最新正式净值口径估算总市值约{total:,.2f}元，其中QDII约{qdii_mv:,.2f}元，境内科技/半导体/通信/云/信息链约{tech_mv:,.2f}元。由于真实现金、014565卖出确认和7月14日010004赎回确认均未完整，报告不计算真实仓位上限，只做条件式行动草案。")
    lines.append("")
    lines.append("| 基金 | 份额 | 最新正式净值日 | 净值 | 估算市值 | 口径 |")
    lines.append("| --- | ---: | --- | ---: | ---: | --- |")
    for r in holdings:
        status = "QDII/港股或海外净值滞后" if r["is_qdii"] == "true" else ("2026-07-14已公布" if r["nav_date"] == "2026-07-14" else "境内正式净值待补")
        lines.append(f"| {r['fund_name']} | {r['shares']} | {r['nav_date']} | {r['nav']} | {float(r['market_value']):,.2f} | {status} |")
    lines.append("")
    lines.append("## 信息面")
    lines.append("")
    lines.append(f"- A股午间：上证{MARKET['sse']}、沪深300 {MARKET['csi300']}、科创50 {MARKET['star50']}、创业板指{MARKET['chinext']}、中证500 {MARKET['csi500']}。成交额：{MARKET['turnover']}。")
    lines.append(f"- 港股已开市：{MARKET['hk']}；隔夜海外：{MARKET['us']}。港股科技上涨与A股科创下跌相背，说明今天不是简单的全球科技同步行情，而是A股半导体/长鑫事件驱动的内部再定价。")
    lines.append("- 国家统计局7月15日10:00发布上半年经济数据日程已触发，但午报抓取时未稳定取得正式解读全文；因此只作为宏观风险偏好节点，不把未核验数字写入因果链。")
    lines.append("- 人民银行7月15日开展4265亿元7天期逆回购，利率1.40%，全额满足一级交易商需求。它支持短端流动性预期，但不能单独抵消半导体IPO兑现和科技拥挤出清。")
    lines.append("- 外交部确认2026世界人工智能大会暨人工智能全球治理高级别会议7月17日至20日在上海举行，国家主席习近平将出席开幕式并发表主旨讲话。上海市政府披露大会有140余场论坛、1100余家企业参展、智算和具身两大赛道各超200家企业。事件重要，但领导出席和会议召开不等于AI链普涨，必须等具体政策、订单或资金流验证。")
    lines.append("")
    lines.append("## 技术面")
    lines.append("")
    lines.append("- 正式净值技术文件显示19只基金中4只站上20日均线、15只低于20日均线；数据员同步后仍没有2026-07-15正式净值，因此评分只到2026-07-14。")
    lines.append("- 盘中估算显示持仓分化：")
    for code, (name, est) in ESTIMATES.items():
        lines.append(f"  - {name}：估算{est}（2026-07-15 11:30，估算净值，不是成交价）。")
    lines.append("- 科创50较沪深300落后约3.71个百分点，易方达半导体设备ETF联接C估算较沪深300落后约5.27个百分点，是今天最清晰的持仓压力来源。")
    lines.append("")
    lines.append("## 资金流与主力资金动向")
    lines.append("")
    lines.append("- 资金口径说明：行情软件“主力资金/大单资金”是按成交方向和单笔规模估算的算法口径，不等同于交易所结算资金；全市场、行业和个股口径不可直接相加。北向资金实时净买入口径近年已不适合作为盘中阈值，本报告不用它触发交易。")
    lines.append("- 可观察结论：全市场成交额线性外推低于3万亿元强修复阈值；科创50大幅弱于宽基；半导体成分样本中下跌多于上涨，通信设备样本涨跌接近但偏弱；计算机/软件相对较强。")
    lines.append("- 主力方向判断：不是市场普遍撤离，而是板块轮动和长鑫发行价落地后的半导体链兑现。若下午扩散到通信和云计算，才从“半导体链兑现”升级为“科技主线撤离”。")
    lines.append("")
    lines.append("## 催化逐条检查")
    lines.append("")
    lines.append("| 催化 | 是否新增 | 进展 | 验证信号 | 失效信号 | 下一检查点 | 评价 |")
    lines.append("| --- | --- | --- | --- | --- | --- | --- |")
    lines.append("| 长鑫科技IPO/发行价 | 既有，追加盘中验证 | 发行价8.66元和7月16日申购已进入交易，半导体设备联接估算明显承压 | 14:45仍弱于沪深300超过2个百分点、半导体资金不收窄 | 半导体设备材料强于沪深300至少1个百分点，通信/云计算不跟跌 | 7月15日14:30/14:45，7月16日申购热度 | 事件超预期在于市值锚和体量，但价格已开始反映；反方是低于万亿叙事反而提升上市空间。 |")
    lines.append("| WAIC/国家领导人出席 | 既有 | 7月17日至20日举行；开幕式讲话待落地 | 出现算力、数据、央企应用、财政/订单等具体措施且电子/通信连续净流入 | 讲话只有治理原则、无新增产业措施；科技高开低走 | 7月17日开幕式和讲话全文 | 政策信号重要但不能直接等同普涨；传导链是战略信号->订单/资金->基金NAV，当前后两环仍是推断。 |")
    lines.append("| 国家统计局上半年数据 | 既有 | 7月15日10:00日程触发，午报未稳定取得正式数据全文 | 宽基、消费/制造和主动权益同步走强 | 数据后指数冲高回落，宽基资金净流出 | 7月15日收盘和晚报 | 宏观节点会影响风险偏好，但不能替代持仓行业验证。 |")
    lines.append("| 人民银行4265亿元逆回购 | 新增 | 7月15日09:20官方公告 | 成交额维持、宽基守3900、资金流出收窄 | 流动性投放后科技仍扩散下跌 | 7月15日14:30/14:45 | 流动性偏正面但新意低，不足以授权申购。 |")
    lines.append("| 高盛中国AI转述 | 既有watch | 仍缺可公开原始报告 | 原始报告或客户摘要、外资/ETF连续流入 | 找不到原报告或板块继续弱 | 继续只作辅助线索 | 二手转述不能独立提高置信度。 |")
    lines.append("| 泰国总理访华 | 既有scheduled | 7月15日至18日 | 官方通稿出现贸易、旅游、产业合作清单 | 仅礼节性会见 | 7月15日至18日通稿 | 对当前科技/半导体仓位影响弱。 |")
    lines.append("")
    lines.append("## 未来催化日历")
    lines.append("")
    lines.append("| 日期 | 事件 | 首次发布时间 | 是否已交易 | 组合影响 |")
    lines.append("| --- | --- | --- | --- | --- |")
    lines.append("| 2026-07-15 | 国家统计局上半年国民经济运行数据发布窗口 | 国家统计局日程表 | 上午已部分交易 | 宽基和主动权益风险偏好，不直接触发科技申赎 |")
    lines.append("| 2026-07-15 | 央行4265亿元7天期逆回购 | 2026-07-15 09:20:31 | 已交易 | 流动性托底，不能单独申购 |")
    lines.append("| 2026-07-16 | 长鑫科技网上/网下申购 | 2026-07-14发行价公告 | 今天已提前交易 | 半导体设备材料和电子特气持仓的核心风险日 |")
    lines.append("| 2026-07-17 | WAIC开幕式及主旨讲话 | 外交部2026-07-13 | 部分预期已交易 | AI/算力/通信/云计算看具体政策和资金确认 |")
    lines.append("| 2026-07-20 | 长鑫缴款压力观察 | 发行安排 | 尚未完全交易 | 观察科技主线资金是否回流或继续虹吸 |")
    lines.append("| 2026-07-15至18 | 泰国总理访华 | 外交部外事日程 | 尚未充分交易 | 低相关，只有具体经贸清单才影响宽基/新兴市场 |")
    lines.append("")
    lines.append("## 午后与1-3个交易日情景")
    lines.append("")
    lines.append("| 情景 | 概率 | 验证信号 | 失效条件 | 对持仓影响 | 行动草案 |")
    lines.append("| --- | ---: | --- | --- | --- | --- |")
    lines.append("| 基准：宽基守稳、科技内部分化 | 45% | 上证守3900，成交外推2.2-2.7万亿元，半导体弱但通信/云未同步崩 | 上涨家数跌破1000或电子流出扩大 | 007339稳定，021894/010004/006503承压 | 等待；只保留021894条件式33%赎回 |")
    lines.append("| 偏强：政策预期与资金回补 | 25% | 半导体设备材料强于沪深300至少1个百分点，通信/云转强，14:45接近日内高位 | 资金流仍负或从高点回落>2个百分点 | 科技估算净值修复 | 不赎回，不申购，等正式净值和WAIC细节 |")
    lines.append("| 偏弱：长鑫虹吸扩散 | 30% | 021894/半导体设备材料弱于沪深300>2个百分点，通信弱于>1.5个百分点，上涨家数<1000或电子资金继续扩大流出 | 上证收复并资金回补 | 半导体、通信、信息行业一起拖累组合 | 截止前赎回021894 50%，再审查007818 20% |")
    lines.append("")
    lines.append("## 对持仓影响")
    lines.append("")
    lines.append("- 易方达沪深300ETF联接C：宽基信号中性，暂不动。")
    lines.append("- 易方达半导体设备ETF联接C：今天最敏感，长鑫从“产业催化”转为“发行价兑现/资金腾挪验证”。若14:45不修复，优先作为结构性减仓对象。")
    lines.append("- 景顺长城电子信息产业股票C：电子特气/半导体材料链暴露，7月14日已有20%赎回申请；今天先查平台是否受理，不重复提交。")
    lines.append("- 国泰中证全指通信设备ETF联接C和华泰柏瑞中证沪港深云计算产业ETF联接C：只有通信/云计算同步跑输才升级，不因半导体单独下跌机械赎回。")
    lines.append("- QDII组合：正式净值日期滞后，不能用今天A股盘中波动直接决定赎回。")
    lines.append("")
    lines.append("## 睡醒后的14:30—15:00操作卡")
    lines.append("")
    lines.append("- 14:30初检：看半导体设备材料相对沪深300、通信相对沪深300、上证3900、上涨家数、成交额进度。")
    lines.append("- 14:45终检：只允许以下三档，不等15:00收盘后再假设按当日净值成交。平台实际截止时间以销售平台显示为准，默认15:00。")
    lines.append("")
    lines.append("| 情景 | 精确阈值 | 动作 | 对象与份额 | 净值日/费用/时滞 | 推翻条件 |")
    lines.append("| --- | --- | --- | --- | --- | --- |")
    lines.append("| 绿色等待 | 半导体设备材料强于沪深300至少1个百分点，或弱势收窄至1个百分点内；通信/云未跑输；上证>3900 | 等待下一开放日 | 0 | 无交易 | 14:45再次转弱 |")
    lines.append("| 黄色结构降半导体 | 半导体设备材料弱于沪深300>2个百分点，通信弱势未扩散，成交额外推<3万亿元 | 截止前提交赎回 | 易方达半导体设备ETF联接C约33%，模型约548.55份 | 预计2026-07-15正式净值；按0.5%估算费约8.09元，T+1确认/T+2或平台规则到账；官方费率未核验则以平台预览为准 | 半导体14:45收复到弱于沪深300<1个百分点或平台费率/份额冲突 |")
    lines.append("| 橙色扩散 | 半导体弱于沪深300>2个百分点且通信弱于>1.5个百分点，或上涨家数<1000 | 截止前提交赎回 | 易方达半导体设备ETF联接C约50%，模型约831.13份；国泰中证全指通信设备ETF联接C再审查20%，模型约96.05份 | 同上；007818费率未核验，平台预览异常则不做第二笔 | 通信收复、云计算稳定、上证守3900且资金流出收窄 |")
    lines.append("| 红色破位 | 上证<3900且14:45不能收复、上涨家数<1000、电子主力净流出>300亿元 | 截止前提交赎回 | 021894 50%；007818 20%-30%；019024 15%-20%；010004只在昨日赎回未受理时补10%-15% | 普通境内基金预计2026-07-15正式净值；费用和确认以平台为准 | 任一关键阈值不成立或平台状态冲突 |")
    lines.append("")
    lines.append("## 行动草案与费用时滞")
    lines.append("")
    lines.append("申购金额为0元。所有赎回均按份额或比例提交，不按盘中价格成交。021894和007818费率目前主要来自模型/用户声明口径，官方费率未核验，若平台预览费率高于0.5%、显示不可赎回、份额不一致或截止已过，一律等待下一开放日。")
    lines.append("")
    lines.append("## 强制自我审计")
    lines.append("")
    lines.append("- 已到评估日的7月14日午报/复核/终检记录已复核并在decision-journal中标为pending_formal_nav_or_confirmation：2026-07-14正式净值可得，010004净值2.1512；但2026-07-15正式净值、平台确认、真实费用、到账均未可得，无法比较完整费用后结果。")
    lines.append("- 方向复核：7月14日终检选择在强修复但资金未完全确认时减010004 20%，当天正式净值较7月13日小幅修复，说明“当日不是崩盘撤退”判断命中；但作为结构性降仓，它的价值要等后续1-3日与不操作对比。")
    lines.append("- 相对不操作价值：现阶段只能记录假设，不能宣称创造收益。错误类型暂列data_quality，因为平台确认和实际费用缺失。样本不足review-loop门槛，不修改永久参数。")
    lines.append("")
    lines.append("## 最强反方观点")
    lines.append("")
    lines.append("半导体设备材料早盘大跌可能只是长鑫发行价落地后的短线兑现，且港股科技上涨、央行流动性投放和WAIC前政策预期仍在；如果下午资金回补，今天赎回021894会把短波止损变成低位卖出。这个反方成立的证据是14:45半导体相对沪深300明显收窄、通信/云不跟跌、成交额没有缩量且收在日内高位附近。")
    lines.append("")
    lines.append("## 投资知识点：事件催化不等于同向普涨")
    lines.append("")
    lines.append("同一个事件可能同时带来产业利好和资金虹吸。检查顺序是：事件是否新增、超预期在哪里、订单/政策如何传导到基金重仓链条、受益细分是谁、价格是否已提前反映、反方解释能否解释盘面。今天长鑫既支持国产存储产业链长期叙事，也可能短期从其他半导体设备、电子特气和通信方向抽走风险预算；所以动作不是“看见长鑫就买半导体”，而是等14:45看资金和相对强弱。")
    lines.append("")
    lines.append("## 来源与抓取时间")
    lines.extend(f"- {src}" for src in SOURCES)
    lines.append("")
    return "\n".join(lines)


def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.name = "Source Han Sans CN"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Source Han Sans CN")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Source Han Sans CN")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Source Han Sans CN")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 12.5, True, "1F4D78")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r, 10.2)


def set_cell(cell, text: str, bold=False, fill: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, 8.2, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, True, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Source Han Sans CN"
    doc.styles["Normal"].font.size = Pt(10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"{REPORT_DATE} 基金午间研究报告")
    set_font(run, 21, True, "0B2545")
    add_para(doc, f"生成时间：{NOW}；数据截止：{DATA_CUTOFF}；未知价原则：盘中数据只作是否提交申请的信号。")

    table_block: list[str] | None = None
    for line in markdown.splitlines()[4:]:
        if not line:
            if table_block:
                flush_table(doc, table_block)
                table_block = None
            continue
        if line.startswith("|"):
            table_block = (table_block or []) + [line]
            continue
        if table_block:
            flush_table(doc, table_block)
            table_block = None
        if line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line[2:])
            set_font(r, 9.8)
        elif line[0:3].count(".") == 1 and line[:1].isdigit():
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line.split(". ", 1)[1] if ". " in line else line)
            set_font(r, 9.8)
        elif line.startswith("# "):
            continue
        else:
            add_para(doc, line)
    if table_block:
        flush_table(doc, table_block)
    doc.save(DOCX)


def flush_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 3:
        return
    headers = [x.strip() for x in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        rows.append([x.strip().replace("<br>", "\n") for x in line.strip("|").split("|")])
    add_table(doc, headers, rows)


def structural_check() -> str:
    with ZipFile(DOCX) as zf:
        names = set(zf.namelist())
        assert "word/document.xml" in names
        xml = zf.read("word/document.xml")
        assert b"2026-07-15" in xml
    reopened = Document(DOCX)
    assert len(reopened.paragraphs) > 30
    assert len(reopened.tables) >= 5
    return f"paragraphs={len(reopened.paragraphs)}, tables={len(reopened.tables)}"


def main() -> None:
    append_catalysts()
    update_decision_journal()
    markdown = make_markdown()
    MD.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(structural_check())
    print(DOCX)
    print(MD)


if __name__ == "__main__":
    main()
