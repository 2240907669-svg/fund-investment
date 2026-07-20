from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Hiragino Sans GB"
INK = "14212B"
BLUE = "1F4D78"
MUTED = "5E6B75"
GRID = "B9C3CC"
HEADER_FILL = "E9EEF3"
TABLE_WIDTH = 9360


def set_run(run, size: float, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().rFonts
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=5, line=1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text: str, size=10.5, bold=False, color=INK) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        set_run(paragraph.add_run(clean), size, bold or is_bold, color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run(run, 8, False, MUTED)


def add_furniture(section) -> None:
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(hp, after=0, line=1.0)
    set_run(hp.add_run("基金午间研究报告 | 仅供个人研究"), 8, False, MUTED)
    add_page_number(section.footer.paragraphs[0])


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def column_widths(count: int) -> list[int]:
    patterns = {
        7: [3100, 850, 1150, 800, 950, 1250, 1260],
        6: [1350, 2050, 950, 1350, 2300, 1360],
        5: [1800, 700, 2200, 2200, 2460],
        4: [1000, 2500, 1400, 4460],
    }
    if count in patterns:
        return patterns[count]
    base = TABLE_WIDTH // count
    widths = [base] * count
    widths[-1] += TABLE_WIDTH - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)


def set_cell(cell, text: str, header: bool, font_size: float) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header or len(text) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, after=0, line=1.02)
    add_text(p, text, font_size, header, INK)
    tc_pr = cell._tc.get_or_add_tcPr()
    if header:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), HEADER_FILL)
        tc_pr.append(shd)
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [[value.strip() for value in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) < 2:
        return
    data = [rows[0]] + rows[2:]
    cols = len(data[0])
    table = doc.add_table(rows=1, cols=cols)
    font_size = 6.7 if cols >= 6 else 7.5
    for idx, value in enumerate(data[0]):
        set_cell(table.rows[0].cells[idx], value, True, font_size)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for source in data[1:]:
        row = table.add_row()
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for idx, value in enumerate(source):
            set_cell(row.cells[idx], value, False, font_size)
    set_table_geometry(table, column_widths(cols))
    spacer = doc.add_paragraph()
    set_spacing(spacer, after=2, line=1.0)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_spacing(p, before=0, after=8, line=1.0)
        add_text(p, text, 20, True, "0B2545")
    elif level == 2:
        set_spacing(p, before=12, after=6, line=1.0)
        add_text(p, text, 14, True, BLUE)
        p.paragraph_format.keep_with_next = True
    else:
        set_spacing(p, before=8, after=4, line=1.0)
        add_text(p, text, 11.5, True, BLUE)
        p.paragraph_format.keep_with_next = True


def build(source: Path, output: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_furniture(doc.sections[0])

    table_lines: list[str] = []
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line.startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []
        if not line:
            continue
        if line == "<!-- PAGEBREAK -->":
            doc.add_page_break()
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text(p, line[2:], 10.2)
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_text(p, re.sub(r"^\d+\.\s", "", line), 10.2)
        else:
            p = doc.add_paragraph()
            set_spacing(p)
            add_text(p, line)
    if table_lines:
        add_table(doc, table_lines)

    props = doc.core_properties
    props.title = output.stem
    props.subject = "场外开放式基金午间研究与下午行动卡"
    props.author = "Codex"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    reopened = Document(output)
    text = "\n".join(p.text for p in reopened.paragraphs)
    if "只看这一页：行动结论" not in text or len(reopened.tables) < 4:
        raise RuntimeError("DOCX structural check failed")
    print(f"created={output} paragraphs={len(reopened.paragraphs)} tables={len(reopened.tables)}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_noon_docx_from_md.py input.md output.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
