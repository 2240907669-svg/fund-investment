from __future__ import annotations

import csv
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "2026-07-14"
NOW = "2026-07-14T20:34:00+08:00"
DOCX = ROOT / "reports" / f"{REPORT_DATE}-基金晚报.docx"
MD = ROOT / "reports" / f"{REPORT_DATE}-基金晚报.md"
DECISION_ID = "DEC-20260714-EVENING-001"


def read_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def set_run_font(run, size=None, bold=None, color=None, name="Source Han Sans CN"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Source Han Sans CN")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_run_font(r, size=8.6, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 11}.get(level, 11), bold=True, color="1F4D78" if level < 3 else "333333")
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.2, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=9.8)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, fill="E8EEF5")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        table_row = table.add_row()
        set_row_cant_split(table_row)
        cells = table_row.cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def latest_navs():
    latest = {}
    for row in read_csv("data/nav-history.csv"):
        if row.get("is_estimate") == "true":
            continue
        code = row["fund_code"]
        key = (row["date"], row["fetched_at"])
        if code not in latest or key > (latest[code]["date"], latest[code]["fetched_at"]):
            latest[code] = row
    return latest


def build_data():
    holdings = read_csv("data/holdings.csv")
    funds = {row["fund_code"]: row for row in read_csv("data/funds.csv")}
    intake = {row["fund_code"]: row for row in read_csv("data/holding-intake.csv")}
    tech = {row["fund_code"]: row for row in read_csv(f"reports/{REPORT_DATE}-holding-technical-indicators.csv")}
    nav = latest_navs()
    rows = []
    total = 0.0
    clusters = defaultdict(float)
    qdii_total = 0.0
    for h in holdings:
        code = h["fund_code"]
        f = funds.get(code, {})
        n = nav.get(code, {})
        shares = float(h["shares"] or 0)
        latest_date = n.get("date", "")
        latest_nav = float(n.get("nav", 0) or 0)
        mv = shares * latest_nav
        total += mv
        cluster = intake.get(code, {}).get("cluster", f.get("theme", ""))
        clusters[cluster] += mv
        if f.get("is_qdii") == "true":
            qdii_total += mv
        rows.append({
            "code": code,
            "name": f.get("fund_name", code),
            "shares": shares,
            "nav_date": latest_date,
            "nav": latest_nav,
            "mv": mv,
            "theme": cluster,
            "is_qdii": f.get("is_qdii") == "true",
            "tech": tech.get(code, {}),
        })
    return rows, total, clusters, qdii_total


def append_decision_journal():
    path = ROOT / "data/decision-journal.csv"
    with path.open(encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f))
        fieldnames = rows[0].keys()
    if any(r["decision_id"] == DECISION_ID for r in rows):
        return
    new = {k: "" for k in fieldnames}
    new.update({
        "decision_id": DECISION_ID,
        "record_origin": "native_realtime",
        "report_time": NOW,
        "data_cutoff": NOW,
        "decision_for_date": "2026-07-15",
        "horizon_days": "10",
        "scope": "portfolio",
        "original_action": "no subscription; wait for 010004 submitted redemption confirmation; next open day only conditional structural review",
        "target_funds": "010004;019024;006503;021894;007818;007339",
        "position_size": "010004 20% already submitted on 2026-07-14 if platform accepted; next-day incremental redemption 0% unless triggers fire",
        "up_probability": "0.40",
        "flat_probability": "0.35",
        "down_probability": "0.25",
        "thesis": "A-share close confirmed broad repair but semiconductor main-fund outflow and incomplete NAV disclosure keep tech holdings in review rather than add mode",
        "trigger_conditions": "SSE above 3900; advancers above 3500; electronic/communication main flow non-negative; total turnover above 2.6tn CNY; formal NAV and platform fee status confirmed",
        "invalidation_conditions": "SSE loses 3900 with advancers below 1000 or semiconductor/electronic outflow expands above 30bn CNY; WAIC speech has no concrete policy and tech relative weakness resumes",
        "execution_deadline": "2026-07-15T15:00:00+08:00",
        "expected_nav_date": "2026-07-15 for ordinary domestic OTC funds submitted before cutoff; QDII NAV lags by product calendar",
        "evaluation_due_date": "2026-07-22",
        "evaluation_status": "pending",
        "evidence_sources": "reports/2026-07-14-基金晚报.docx;data/catalyst-ledger.csv;data/nav-history.csv;reports/2026-07-14-holding-technical-indicators.csv;https://qt.gtimg.cn/;https://push2.eastmoney.com/;https://www.pbc.gov.cn/;https://www.stats.gov.cn/sj/fbrc/bnxxfb/;https://www.fmprc.gov.cn/web/wjdt_674879/wsrc_674883/;https://www.shanghai.gov.cn/nw4411/20260708/ba4c8e75f2744b43a6080ebb82a3aab2.html",
    })
    rows.append(new)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def build_docx():
    holdings, total, clusters, qdii_total = build_data()
    fund_name = {h["code"]: h["name"] for h in holdings}
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Source Han Sans CN"
    styles["Normal"]._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Source Han Sans CN")
    styles["Normal"]._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Source Han Sans CN")
    styles["Normal"]._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Source Han Sans CN")
    styles["Normal"].font.size = Pt(10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("2026-07-14 基金晚间决策型复盘")
    set_run_font(r, size=22, bold=True, color="0B2545")
    add_para(doc, f"生成时间：{NOW}；账户范围：大陆场外开放式基金；盘中估值仅作信号，不是成交价。", bold=True)
    add_para(doc, "结论先行：今日 A 股收盘修复强于午盘，但半导体资金与多数持仓技术结构仍未完全确认。已提交的景顺长城电子信息产业股票C 20% 赎回等待正式净值与平台确认；下一开放日不新增申购，不做账户级撤退。")

    add_heading(doc, "只看这一页：行动结论", 1)
    add_table(doc, ["问题", "结论"], [
        ["1）下一开放日卖不卖", "继续持有：易方达沪深300ETF联接C、易方达半导体材料设备ETF联接C、国泰中证全指通信设备ETF联接C、华泰柏瑞中证沪港深云计算产业ETF联接C和QDII组合，等待正式净值和7月17日世界人工智能大会。优先减仓审查：景顺长城电子信息产业股票C已于7月14日提交约20%赎回（模型值432.92份）待确认；它按电子特气/半导体材料链暴露审查，不按泛电子行业简单处理。若未受理且7月15日电子特气、半导体材料或半导体资金再度转弱，可重新提交10%-15%。易方达信息行业精选股票C仅在电子/计算机资金重新转负且跌回弱势时审查15%-20%（模型值89.60-119.46份，按最新正式净值估算毛额约270.70-361.02元，费率0.5%口径）。触发即退出：暂无账户级退出项；财通集成电路产业股票C金额小且费率/状态未知，暂不动作。"],
        ["2）下一开放日买不买", "不买。申购金额上限为0元。只有在现金真实余额、申购状态、费率确认，且NBS数据后宽基成交延续、电子/通信连续净流入、WAIC出现具体政策或订单时，才重新评估分批买入；本晚报不授权预埋申购。"],
        ["3）跑不跑", "不需要账户级跑。当前是结构性降仓而非清仓：景顺长城电子信息产业股票C已做电子特气/半导体材料链弱势仓位的风险释放，其余等待确认。若账户高点回撤经真实流水确认达到8%，再切换到风险资产上限40%的纪律模式。"],
        ["4）概率", "下一交易日组合：收涨40%、震荡35%、收跌25%。未来1-2周偏强概率45%，震荡35%，偏弱20%。"],
        ["5）明天只做一件事", "10:00看国家统计局数据后，到14:45只复核五个阈值；不在没有平台确认的情况下追加操作。"],
        ["6）最多五个指标", "沪指3900；上涨家数3500/1000；电子或半导体主力资金0/-300亿元；通信主力资金是否继续为正；沪深京成交额是否维持2.6万亿元以上。"],
        ["7）推翻结论的证据", "沪指跌回3900下、上涨家数少于1000、电子特气/半导体材料链和电子信息资金流出扩大，且景顺长城电子信息产业股票C、易方达信息行业精选股票C正式净值继续弱于沪深300。"],
    ], [3.4, 13.0])

    add_heading(doc, "持仓复述与正式净值", 1)
    add_para(doc, f"同步后持仓按最新正式净值估算总市值约 {total:,.2f} 元。项目自检仍提示：组合历史为空，真实账户回撤无法计算；持仓市值与配置资金 30,000 元差异超过10%，现金比例不能直接采用配置值。")
    nav_rows = []
    for h in holdings:
        status = "待公布" if h["nav_date"] != REPORT_DATE else "已公布"
        if h["is_qdii"]:
            status = f"QDII净值日期滞后：{h['nav_date']}"
        elif h["nav_date"] != REPORT_DATE:
            status = f"{REPORT_DATE}待公布；最新{h['nav_date']}"
        nav_rows.append([h["name"], f"{h['shares']:.2f}", h["nav_date"], f"{h['nav']:.4f}", f"{h['mv']:,.2f}", status])
    add_table(doc, ["基金", "份额", "净值日", "净值", "估算市值", "披露状态"], nav_rows, [4.9, 1.8, 1.8, 1.5, 2.0, 4.2])

    add_heading(doc, "广泛信息面", 1)
    add_bullet(doc, "A股收盘：上证指数3967.13（+1.36%），沪深300 4796.50（+2.15%），创业板指3851.14（+3.43%），科创50 2009.73（+0.77%）；沪深两市成交额约2.70万亿元。")
    add_bullet(doc, "市场宽度：上证和深证成分统计合计上涨约3967家、下跌约1234家，满足强修复的广度条件，但不是所有主题同步修复。")
    add_bullet(doc, "港股：恒生指数24340.73（+0.52%），恒生科技指数4679.46（+0.06%），对A股科技修复的确认偏弱。")
    add_bullet(doc, "海外：前一美股交易日纳指-1.55%、标普500-0.79%，全球科技链仍是反方压力。")
    add_bullet(doc, "央行：7月14日9:20公告开展2365亿元7天期逆回购，中标利率1.40%，偏流动性呵护，但不能单独解释行业上涨。")

    add_heading(doc, "未来催化日历", 1)
    add_table(doc, ["日期", "事件", "新增/复核", "传导链与反方", "下一检查点"], [
        ["7/15 10:00", "国家统计局发布国民经济运行情况", "复核：官方日程确认", "宏观数据->宽基风险偏好；反方是数据后冲高回落或政策预期落空", "14:45看沪深300、中证500、成交额"],
        ["7/15-18", "泰国总理访华", "新增：外交部外事日程", "外交->经贸/旅游/区域合作；反方是礼节性会见，无订单/政策细则", "会见通稿是否有合作清单"],
        ["7/16起", "长鑫科技科创板申购", "复核：继续scheduled", "国产存储->设备材料/封测；反方是大IPO分流和估值已定价", "申购热度、定价、半导体资金"],
        ["7/17-20", "WAIC暨AI全球治理高级别会议", "复核：外交部确认最高层出席，上海市政府确认论坛与展品", "AI战略->算力/芯片/通信/云；反方是大会已预告，若无具体政策可能兑现", "7/17讲话全文和电子/通信连续资金"],
        ["7/14", "央行2365亿元7天逆回购", "新增：官方公告", "流动性->成交和风险偏好；反方是例行操作", "7/15公开市场与成交延续"],
    ], [2.1, 3.1, 2.5, 6.0, 3.0])

    add_heading(doc, "技术面", 1)
    tech_rows = []
    focus = ["007339", "021894", "010004", "007818", "019024", "006503", "012922", "006479", "019331"]
    for code in focus:
        item = next((h for h in holdings if h["code"] == code), None)
        if not item:
            continue
        t = item["tech"]
        tech_rows.append([item["name"], t.get("latest_date", ""), t.get("nav_vs_ma20", ""), t.get("nav_vs_ma60", ""), t.get("rsi14", ""), t.get("drawdown_60obs", ""), t.get("relative_20obs_vs_007339", ""), t.get("technical_state", "")])
    add_table(doc, ["基金", "净值日", "相对MA20", "相对MA60", "RSI14", "60日回撤", "20日相对沪深300", "状态"], tech_rows, [4.3, 1.6, 1.6, 1.6, 1.2, 1.5, 1.8, 2.6])
    add_para(doc, "技术结论：易方达半导体材料设备ETF联接C、华泰柏瑞中证沪港深云计算产业ETF联接C仍有相对强弱支撑；景顺长城电子信息产业股票C、易方达信息行业精选股票C、国泰中证全指通信设备ETF联接C和多数主动科技仓仍在20日线下。景顺长城电子信息产业股票C要按电子特气/半导体材料链验证，今天的价格修复尚未等于趋势恢复。QDII多只净值滞后，不能拿A股或美股盘中估算替代正式净值。")

    add_heading(doc, "主力资金动向", 1)
    add_para(doc, "口径说明：所谓主力资金是行情软件基于逐笔成交方向的统计，不等同于机构真实账户流向；必须与价格、成交额、宽度和相对强弱交叉验证。")
    add_table(doc, ["方向", "行业", "净额", "价格验证", "解释"], [
        ["流入", "元件", "+141.4亿元", "+6.31%", "AI硬件/PCB链条修复最强"],
        ["流入", "印制电路板", "+132.8亿元", "+6.88%", "与WAIC/算力硬件预期一致，但需看持续性"],
        ["流入", "通信设备+通信", "约+174.4亿元", "+1.59%/+0.82%", "007818相关方向改善，支持继续观察而非追买"],
        ["流出", "半导体", "-124.2亿元", "+0.17%", "价格止跌但资金未确认，易方达半导体材料设备ETF联接C不减但不加；景顺长城电子信息产业股票C也受电子特气/半导体材料链资金验证约束"],
        ["流出", "计算机", "-49.3亿元", "+0.08%", "AI软件/应用端确认不足"],
    ], [1.4, 3.2, 2.0, 2.2, 7.5])
    add_para(doc, "判断：全市场不是集体出逃，而是强修复中的板块轮动。资金从部分半导体、计算机、军工电子流出，转向PCB、元件、通信和有色。对持仓的含义是：不追买科技，但也不把今日修复解读成账户级撤退信号。")

    add_heading(doc, "情景预测与行动", 1)
    add_table(doc, ["周期", "情景", "概率", "因果链", "验证信号", "行动", "失效条件"], [
        ["1日", "基准：震荡偏强", "45%", "宏观数据前风险偏好修复，科技内部继续分化", "沪指>3900、成交>2.6万亿、通信为正", "不买；景顺长城电子信息产业股票C确认后观察", "电子特气/半导体材料链流出扩大"],
        ["1日", "偏强", "30%", "NBS数据不差+WAIC预期升温+资金扩散", "上涨>3500且电子/通信均为正", "只取消新增减仓，不追申购", "冲高回落且宽度收缩"],
        ["1日", "偏弱", "25%", "海外科技压力或政策预期落空，半导体再流出", "沪指<3900、上涨<1000、电子特气/半导体材料链相关资金流出扩大", "重新审查易方达信息行业精选股票C和景顺长城电子信息产业股票C 10%-20%", "沪指收复3900且流出收窄"],
        ["1-2周", "基准：震荡修复", "35%", "大会和宏观数据提供事件窗口，但正式政策仍未落地", "宽基MA20修复、科技相对强弱不再恶化", "维持仓位，先补数据", "连续两日跌破关键位"],
        ["1-2周", "偏强", "45%", "WAIC出现具体算力/数据/应用政策或订单，成交延续", "电子特气/半导体材料链和通信连续净流入，易方达半导体材料设备ETF联接C、华泰柏瑞中证沪港深云计算产业ETF联接C强于沪深300", "再评估小额分批，当前不预埋", "讲话无新增细则"],
        ["1-2周", "偏弱", "20%", "事件已定价，技术修复失败", "景顺长城电子信息产业股票C和易方达信息行业精选股票C正式净值继续弱于宽基", "结构性降仓，不账户级清仓", "宽基放量突破且资金扩散"],
    ], [1.2, 2.3, 1.1, 4.0, 3.3, 3.0, 3.0])

    add_heading(doc, "对持仓的影响", 1)
    add_bullet(doc, "宽基：易方达沪深300ETF联接C收盘确认修复，但7月14日正式净值待公布；继续持有，作为组合稳定器，不因盘中点位做申赎。")
    add_bullet(doc, "半导体材料与电子特气：易方达半导体材料设备ETF联接C技术相对最强但半导体资金未确认，继续持有不加仓；景顺长城电子信息产业股票C已提交20%结构性赎回，等待正式净值和平台确认。该基金不是泛“电子”处理，按广钢气体等电子特气/半导体材料链暴露审查；易方达信息行业精选股票C仅保留触发式审查。")
    add_bullet(doc, "通信/云：通信资金与价格相互验证，国泰中证全指通信设备ETF联接C和华泰柏瑞中证沪港深云计算产业ETF联接C不减；但国泰中证全指通信设备ETF联接C正式净值待公布，不能按指数涨幅估算成交。")
    add_bullet(doc, "主动权益：财通成长优选混合C、财通品质甄选混合C今日正式净值修复，但多数仍低于MA20，不加仓；先观察明日宏观数据后的宽基延续。")
    add_bullet(doc, "QDII：广发纳指100ETF联接(QDII)C、华夏移动互联灵活配置混合(QDII)、建信新兴市场优选混合(QDII)A、国富全球科技互联混合(QDII)人民币A的净值仍停在7月10日；易方达全球成长精选混合(QDII)C、嘉实全球产业升级股票(QDII)C、天弘全球高端制造混合(QDII)C到7月13日。只做风险观察，不用A股收盘替代其正式净值。")

    add_heading(doc, "午报14:30-15:00操作卡复盘", 1)
    add_para(doc, "午报阈值复核：确认破位三条件没有触发。14:32记录显示沪指3952.10、上涨4029家，强价格修复成立；电子主力资金仍约-41.1亿元，成交投射约2.76万亿元，低于强资金确认阈值，因此触发的是“强修复但资金未确认”的分支。")
    add_para(doc, "用户已提交景顺长城电子信息产业股票C赎回申请（交易账本记录时间：2026-07-14T14:43:59+08:00，状态为submitted）。若平台确在15:00前受理，普通场外基金预计适用2026-07-14未知价正式净值；截至本报告该基金的7月14日正式净值仍待公布，不能判定成交结果、费用后价值或午报胜负。")
    add_para(doc, "只有收盘后才确认的信号：上证3967.13、沪深成交约2.70万亿元、板块资金最终分布。这些只能用于7月15日及之后的开放日决策，不能伪装成可按7月14日收盘净值重新提交。")

    add_heading(doc, "账户8%回撤纪律", 1)
    add_para(doc, "组合历史为空，无法计算真实高点和账户回撤。按项目纪律，若用户补录真实现金、天弘沪深港创新药50ETF联接C确认、景顺长城电子信息产业股票C确认净值和完整流水后，账户从历史高点回撤达到8%，则停止新增申购，风险资产上限降至40%，并进入至少5个有记录交易日的冷静期。当前不能把未知回撤当作0，也不能据此追加买入。")

    add_heading(doc, "审计与滚动指标", 1)
    add_para(doc, "今日午报原始预测仍为pending：正式NAV和平台确认未完成。实际执行、按原建议模型执行、完全不操作和宽基基准四项比较暂不能计算。历史日志样本少，native实时已到7条以内且多条未到评估日；不得宣称胜率提升。")
    add_table(doc, ["指标", "当前值", "说明"], [
        ["可解析日志样本", "8条（含回填）", "回填记录不冒充实时样本"],
        ["已完成评估样本", "0条", "正式净值/执行数据不足"],
        ["方向命中率", "样本不足", "不报告胜率"],
        ["Brier分数", "样本不足", "需要概率和实际结果"],
        ["相对不操作费用后增益", "样本不足", "景顺长城电子信息产业股票C仍待正式净值"],
        ["最大不利波动", "样本不足", "缺少组合历史"],
    ], [3.0, 3.0, 10.0])

    add_heading(doc, "今日教学：资金确认不等于价格上涨", 1)
    add_para(doc, "概念：价格上涨说明有人愿意用更高价格成交，但资金确认要看上涨是否由广度、成交额、相对强弱和主力资金共同支持。若只有少数链条上涨，场外基金不应追逐当天情绪。")
    add_para(doc, "持仓例子：易方达半导体材料设备ETF联接C今日拿到7月14日正式净值2.9478，仍在MA20上方，但半导体行业主力资金净流出约124.2亿元。因此结论是继续持有、等待世界人工智能大会和长鑫科技事件验证，而不是新增申购。景顺长城电子信息产业股票C则要额外看电子特气/半导体材料链是否修复，因为官方一季报显示其第一大重仓是广钢气体。")

    add_heading(doc, "来源与时间戳", 1)
    for src in [
        "项目文件：AGENTS.md、config/trading-constraints.json、config/noon-decision-rules.json、config/review-loop.json、data/catalyst-ledger.csv、data/decision-journal.csv、data/nav-history.csv，抓取/生成时间见各文件。",
        "基金净值：东方财富基金 pingzhongdata，同步时间2026-07-14T12:30:21Z；正式净值只采用is_estimate=false记录。",
        "行情：腾讯行情 https://qt.gtimg.cn/；东方财富板块资金 https://push2.eastmoney.com/，抓取时间约2026-07-14晚间。",
        "央行公开市场公告：https://www.pbc.gov.cn/zhengcehuobisi/125207/125213/125431/125475/2026071409063784437/index.html",
        "国家统计局发布日程：https://www.stats.gov.cn/sj/fbrc/bnxxfb/",
        "外交部外事日程：https://www.mfa.gov.cn/web/wjdt_674879/wsrc_674883/",
        "上海市政府WAIC信息：https://www.shanghai.gov.cn/nw4411/20260708/ba4c8e75f2744b43a6080ebb82a3aab2.html",
        "景顺长城电子信息产业股票2026年第1季度报告：官方季报显示第一大重仓为广钢气体，占基金资产净值9.27%，据此把该基金按电子特气/半导体材料链暴露审查。",
    ]:
        add_bullet(doc, src)

    doc.save(DOCX)

    MD.write_text(
        f"# {REPORT_DATE} 基金晚报\n\nDOCX已生成：{DOCX}\n\n"
        "核心结论：不买；景顺长城电子信息产业股票C 20%已提交待确认；下一开放日只做触发式结构性审查，不账户级撤退。\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    append_decision_journal()
    build_docx()
    print(DOCX)
