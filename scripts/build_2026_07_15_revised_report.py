from __future__ import annotations

import csv
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "reports" / "2026-07-15-基金午报-14时修正版.md"
DOCX = ROOT / "reports" / "2026-07-15-基金午报-14时修正版.docx"
DECISION_ID = "DEC-20260715-REVISED-1408-001"


def set_font(run, size=10.2, bold=False, color="000000"):
    run.font.name = "Source Han Sans CN"
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        fonts.set(qn(key), "Source Han Sans CN")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def clean_text(text):
    return text.replace("**", "")


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(clean_text(text))
    set_font(run, 8.0, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def flush_table(doc, lines):
    if len(lines) < 3:
        return
    headers = [cell.strip() for cell in lines[0].strip("|").split("|")]
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines[2:]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, True, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx():
    markdown = MD.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    doc.styles["Normal"].font.name = "Source Han Sans CN"
    doc.styles["Normal"].font.size = Pt(10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(title.add_run("2026-07-15 基金午报（14时修正版）"), 20, True, "0B2545")

    table_lines = []
    for line in markdown.splitlines()[1:]:
        if line.startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            flush_table(doc, table_lines)
            table_lines = []
        if not line:
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            set_font(p.add_run(clean_text(line[3:])), 13.5, True, "1F4D78")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2.5)
            set_font(p.add_run(clean_text(line[2:])), 9.6)
        elif len(line) > 2 and line[0].isdigit() and ". " in line[:4]:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2.5)
            set_font(p.add_run(clean_text(line.split(". ", 1)[1])), 9.6)
        elif line.startswith("# "):
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3.5)
            p.paragraph_format.line_spacing = 1.08
            set_font(p.add_run(clean_text(line)), 9.8)
    if table_lines:
        flush_table(doc, table_lines)
    doc.save(DOCX)


def append_decision():
    path = ROOT / "data" / "decision-journal.csv"
    with path.open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if any(row["decision_id"] == DECISION_ID for row in rows):
        return
    fieldnames = list(rows[0].keys())
    row = {key: "" for key in fieldnames}
    row.update({
        "decision_id": DECISION_ID,
        "record_origin": "native_realtime",
        "supersedes_decision_id": "DEC-20260715-NOON-001",
        "report_time": "2026-07-15T14:12:00+08:00",
        "data_cutoff": "2026-07-15T14:10:00+08:00",
        "decision_for_date": "2026-07-15",
        "horizon_days": "3",
        "scope": "portfolio",
        "original_action": "wait now; no subscription; redeem 20% of 021894 before cutoff only on joint technology-spread confirmation at 14:45",
        "target_funds": "021894",
        "position_size": "0% now; conditional redemption 20% available shares, modeled 332.45 shares; subscription 0 CNY",
        "up_probability": "0.20",
        "flat_probability": "0.35",
        "down_probability": "0.45",
        "thesis": "Semiconductor equipment materially underperforms CSI300 and weakness has spread to communication and AI, but SSE remains above 3900, advancers exceed decliners, and semiconductor equipment has bounced from the intraday low; wait until joint breakdown confirmation rather than sell from loss aversion or panic.",
        "trigger_conditions": "At 14:45 semiconductor-equipment ETF <=-6.5% and underperforms CSI300 by >=6pp; communication ETF <=-2.5% or AI ETF <=-3%; and semiconductor-equipment ETF within 0.8% of intraday low with negative active-volume imbalance, or market advancers <1800.",
        "invalidation_conditions": "At 14:45 semiconductor-equipment loss narrows inside -5%; communication or AI recovers near CSI300; advancers remain >2500; platform fee, available shares, redemption status, or cutoff conflicts.",
        "execution_deadline": "2026-07-15T15:00:00+08:00; target completion before 14:50; actual platform cutoff prevails",
        "expected_nav_date": "2026-07-15 if accepted before actual platform cutoff; otherwise next open day",
        "evaluation_due_date": "2026-07-20",
        "evaluation_status": "pending",
        "evidence_sources": "reports/2026-07-15-基金午报-14时修正版.docx;https://qt.gtimg.cn/;https://vip.stock.finance.sina.com.cn/quotes_service/api/json_v2.php/Market_Center.getHQNodeData;https://fundgz.1234567.com.cn/;https://cdn.efunds.com.cn/owch/data/bulletin/20250919/;data/nav-history.csv;data/catalyst-ledger.csv",
    })
    rows.append(row)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def structural_check():
    with ZipFile(DOCX) as archive:
        assert "word/document.xml" in archive.namelist()
    reopened = Document(DOCX)
    assert len(reopened.paragraphs) > 50
    assert len(reopened.tables) >= 8
    return len(reopened.paragraphs), len(reopened.tables)


if __name__ == "__main__":
    build_docx()
    append_decision()
    print(structural_check())
    print(DOCX)
